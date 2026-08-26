from docx import Document
from docx.oxml.ns import qn
import json, hashlib, os, sys

def twips(v): return None if v is None else round(v.twips)
def inspect(path):
    d=Document(path)
    out={"path":path,"sha256":hashlib.sha256(open(path,'rb').read()).hexdigest(),"size":os.path.getsize(path),"sections":[],"paragraphs":[],"tables":[]}
    for s in d.sections:
        out["sections"].append({"page_w":twips(s.page_width),"page_h":twips(s.page_height),"top":twips(s.top_margin),"bottom":twips(s.bottom_margin),"left":twips(s.left_margin),"right":twips(s.right_margin),"header":twips(s.header_distance),"footer":twips(s.footer_distance)})
    for i,p in enumerate(d.paragraphs):
        pf=p.paragraph_format
        out["paragraphs"].append({"i":i,"text":p.text,"style":p.style.name,"align":str(p.alignment),"before":twips(pf.space_before),"after":twips(pf.space_after),"line":str(pf.line_spacing),"left":twips(pf.left_indent),"first":twips(pf.first_line_indent),"keep_next":pf.keep_with_next,"runs":[{"text":r.text,"font":r.font.name,"size":twips(r.font.size),"bold":r.bold,"italic":r.italic,"underline":str(r.underline),"color":str(r.font.color.rgb)} for r in p.runs]})
    for ti,t in enumerate(d.tables):
        out["tables"].append({"i":ti,"rows":[[c.text for c in row.cells] for row in t.rows]})
    return out

for p in sys.argv[1:]:
    print(json.dumps(inspect(p),indent=2,ensure_ascii=False))

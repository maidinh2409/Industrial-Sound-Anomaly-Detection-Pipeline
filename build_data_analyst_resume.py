from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"C:\Users\ACER\Industrial-Sound-Anomaly-Detection-Pipeline\Demi_Le_Data_Analyst_Resume.docx"
BLUE = RGBColor(31, 77, 120)
DARK = RGBColor(32, 38, 46)
GRAY = RGBColor(85, 92, 101)

def font(run, size=9.1, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def bottom_border(p, color="A8B6C5", size="8"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), size)
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
    pBdr.append(b)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_margins(cell, top=40, start=90, bottom=40, end=90):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn("w:"+m))
        if node is None: node=OxmlElement("w:"+m); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_repeat_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for width in widths:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(width)); grid.append(gc)
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"),str(width)); tcW.set(qn("w:type"),"dxa")
            set_cell_margins(cell)

def section_heading(doc, text):
    p=doc.add_paragraph(style="Heading 1"); p.paragraph_format.keep_with_next=True
    p.add_run(text.upper()); bottom_border(p)
    return p

def bullet(doc, text):
    p=doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together=True
    p.add_run(text)
    return p

def role(doc, title, org, dates):
    t=doc.add_table(rows=1, cols=2); t.style="Table Grid"
    set_repeat_table_geometry(t,[6900,2460])
    for c in t.rows[0].cells:
        c._tc.get_or_add_tcPr().remove(c._tc.get_or_add_tcPr().find(qn("w:tcBorders"))) if c._tc.get_or_add_tcPr().find(qn("w:tcBorders")) is not None else None
    p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    font(p.add_run(title),9.6,True); font(p.add_run(" | "+org),9.3,False,GRAY)
    p=t.cell(0,1).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after=Pt(0); font(p.add_run(dates),9.1,True,GRAY)

def project(doc, name, stack, date, bullets):
    role(doc,name,stack,date)
    for b in bullets: bullet(doc,b)

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=Inches(.47); sec.bottom_margin=Inches(.47); sec.left_margin=Inches(.62); sec.right_margin=Inches(.62)
sec.header_distance=Inches(.2); sec.footer_distance=Inches(.2)

normal=doc.styles["Normal"]
normal.font.name="Calibri"; normal.font.size=Pt(9.1); normal.font.color.rgb=DARK
normal.paragraph_format.space_after=Pt(2); normal.paragraph_format.line_spacing=1.0
for name,size,before,after in (("Heading 1",11.5,7,3),("Heading 2",10,5,2),("Heading 3",9.5,4,2)):
    s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=BLUE
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
lb=doc.styles["List Bullet"]
lb.font.name="Calibri"; lb.font.size=Pt(9.1); lb.font.color.rgb=DARK
lb.paragraph_format.left_indent=Inches(.23); lb.paragraph_format.first_line_indent=Inches(-.13); lb.paragraph_format.space_after=Pt(1.2); lb.paragraph_format.line_spacing=1.0

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
font(p.add_run("DEMI LE"),20,True,BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
font(p.add_run("DATA ANALYST"),10,True,GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(5)
font(p.add_run("+1 236-234-2409  |  lenguyenmaidinh2409@gmail.com  |  LinkedIn  |  GitHub  |  Portfolio"),8.7,False,GRAY)
bottom_border(p,"2E74B5","10")

section_heading(doc,"Professional Summary")
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
font(p.add_run("Data analyst with experience translating operational, marketing, and large-scale behavioral data into reliable reporting and actionable insights. Skilled in SQL, Python, Excel, Tableau, Power BI, dbt, and BigQuery; experienced in data cleaning, KPI analysis, dashboarding, experimentation, and predictive modeling."),9.1)

section_heading(doc,"Core Skills")
t=doc.add_table(rows=3,cols=2); t.style="Table Grid"; set_repeat_table_geometry(t,[1600,7760])
skills=[("Analytics","SQL (joins, CTEs, subqueries), Python, Excel (XLOOKUP, INDEX/MATCH, dynamic arrays), data cleaning, EDA, KPI reporting, A/B testing"),("BI & Data","Power BI, Tableau, DAX, dbt, Google BigQuery, MySQL, SQL Server, data pipelines, dashboard design"),("Libraries & Methods","Pandas, NumPy, Matplotlib, Seaborn, scikit-learn, hypothesis testing, regression and classification")]
for row,(label,value) in zip(t.rows,skills):
    shade(row.cells[0],"E8EEF5")
    p=row.cells[0].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(label),8.8,True,BLUE)
    p=row.cells[1].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(value),8.7)

section_heading(doc,"Experience")
role(doc,"Test Administrator / Data Support","Pearson Professional Center","Mar 2026 - Present")
bullet(doc,"Automated daily candidate-roster generation with XLOOKUP, INDEX/MATCH, nested logic, and dynamic arrays, reducing manual entry and preventing roster errors.")
bullet(doc,"Extract and prepare daily check-in data while validating schedules and candidate records to maintain accurate operations.")
bullet(doc,"Standardized reporting templates and documented repeatable data workflows for the administration team.")
role(doc,"Marketing Analyst","Heyo Smart Technology","Sep 2023 - May 2024")
bullet(doc,"Analyzed market, competitor, and customer behavior using Google Analytics; translated real-time metrics into KPI reports and campaign decisions.")
bullet(doc,"Ran A/B tests and targeted campaigns that increased website visitors by 50% and generated 150+ clicks in seven days.")
role(doc,"SEO Analyst","EveHR LLC","Jun 2023 - Nov 2023")
bullet(doc,"Analyzed keyword and market trends with SEMrush and Ahrefs, informing SEO, social, and email investment decisions.")
bullet(doc,"Increased website traffic 70% in six months and achieved 20+ top-five keyword rankings through data-driven optimization.")

section_heading(doc,"Selected Analytics Projects")
project(doc,"Exoplanet Habitability Analysis","SQL, Python, Power BI, dbt, BigQuery","May 2025",[
    "Built a cloud analytics pipeline from API extraction and Google Cloud Storage through tested dbt transformations in BigQuery.",
    "Designed an interactive Power BI dashboard to explore planetary characteristics and habitability trends."
])
project(doc,"Gym Check-ins & User Analysis","MySQL, Tableau","Dec 2024",[
    "Queried 300K+ records across five related tables using joins, CTEs, subqueries, and stored functions; visualized usage and customer patterns in Tableau.",
    "Identified Gen Y as the most active segment, with particularly high female engagement, supporting targeted marketing opportunities."
])
project(doc,"Real Estate Sales Analysis","Python, Pandas, Visualization","Jan 2025",[
    "Cleaned and analyzed 1M+ property-sale records, resolving missing values, inconsistencies, and outliers before exploratory analysis.",
    "Found that homes sold for roughly 1.6x assessed value and that 75% took more than one year to sell after listing."
])
project(doc,"Crash Reporting Dashboard","Tableau, Power BI, DAX","Nov 2024",[
    "Collaborated on dashboards for 18K+ collision records using time-series and geographic analysis, calculated fields, table calculations, and DAX."
])

section_heading(doc,"Education")
role(doc,"Diploma, Computer Studies & Information Systems (Data Analytics)","Douglas College","Expected Sep 2026")
role(doc,"Bachelor of Business Management with Marketing","University of the West of England","Sep 2023")
p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.08); p.paragraph_format.space_after=Pt(0); font(p.add_run("Leadership: Vice President, Flagship Startup Club"),8.8,False,GRAY,True)

doc.core_properties.title="Demi Le - Data Analyst Resume"
doc.core_properties.subject="Resume tailored for data analyst roles"
doc.core_properties.author="Demi Le"
doc.save(OUT)
print(OUT)

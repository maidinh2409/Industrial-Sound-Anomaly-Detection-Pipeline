from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import hashlib

REFERENCE=r"D:\Download here\Data Analyst Template.docx"
OUTPUT=r"C:\Users\ACER\Industrial-Sound-Anomaly-Detection-Pipeline\Demi_Le_Data_Analyst_Resume_Template_Format.docx"
EXPECTED="048afabb2bbdf17854ad0bc1a56b491cda4c82da26a1639b57d9eca621d3014f"
assert hashlib.sha256(open(REFERENCE,'rb').read()).hexdigest()==EXPECTED

doc=Document(REFERENCE)
bullet_numpr=None
for source_p in doc.paragraphs:
    numpr=source_p._p.pPr.find(qn('w:numPr')) if source_p._p.pPr is not None else None
    if numpr is not None:
        bullet_numpr=deepcopy(numpr)
        break
body=doc._element.body
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

def set_font(run,size=10,bold=None,italic=None):
    run.font.name='Times New Roman'
    rpr=run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:ascii'),'Times New Roman'); rpr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
    run.font.size=Pt(size); run.font.color.rgb=RGBColor(0,0,0)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic

def base(p,left=-.625,after=0,before=0,line=1.0):
    pf=p.paragraph_format; pf.left_indent=Inches(left); pf.right_indent=Inches(0)
    pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line

def border_bottom(p):
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'1'); b.set(qn('w:color'),'666666')
    pBdr.append(b); pPr.append(pBdr)

def heading(text):
    p=doc.add_paragraph(); base(p,after=1,before=6); p.paragraph_format.keep_with_next=True
    set_font(p.add_run(text),10,bold=True); border_bottom(p)

def textline(text,italic=False,after=0,left=-.5):
    p=doc.add_paragraph(); base(p,left,after=after); set_font(p.add_run(text),10,italic=italic); return p

def entry(left,date,after=0):
    p=doc.add_paragraph(); base(p,-.5,after=after); p.paragraph_format.keep_with_next=True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5),WD_TAB_ALIGNMENT.RIGHT,WD_TAB_LEADER.SPACES)
    if ' | ' in left:
        a,b=left.split(' | ',1)
        set_font(p.add_run(a),10,bold=True)
        set_font(p.add_run(' | '),10)
        set_font(p.add_run(b),10,italic=True)
    else:
        set_font(p.add_run(left),10,bold=True)
    set_font(p.add_run('\t'+date),10)
    return p

def bullet(text):
    p=doc.add_paragraph(); base(p,-.12,after=0)
    p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.first_line_indent=Inches(-.18)
    if bullet_numpr is not None:
        p._p.get_or_add_pPr().append(deepcopy(bullet_numpr))
    set_font(p.add_run(text),10); return p

def skill(label,value):
    p=doc.add_paragraph(); base(p,-.5,after=0)
    set_font(p.add_run(label),10,bold=True); set_font(p.add_run(value),10)

p=doc.add_paragraph(); base(p,-.625); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('DEMI LE'),18,bold=True)
p=doc.add_paragraph(); base(p,-.625); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('DATA ANALYST'),10,bold=True)
p=doc.add_paragraph(); base(p,-.625,after=2); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('+1 236-234-2409  |  lenguyenmaidinh2409@gmail.com  |  LinkedIn  |  GitHub  |  Portfolio'),10)

heading('PROFESSIONAL SUMMARY')
textline('Data analyst with experience translating operational, marketing, and large-scale behavioral data into reliable reporting and actionable insights. Skilled in SQL, Python, Excel, Tableau, Power BI, dbt, and BigQuery; experienced in data cleaning, KPI analysis, dashboarding, experimentation, and predictive modeling.',after=1)

heading('CORE SKILLS')
skill('Analytics ','SQL (joins, CTEs, subqueries), Python, Excel (XLOOKUP, INDEX/MATCH, dynamic arrays), data cleaning, EDA, KPI reporting, A/B testing')
skill('BI & Data ','Power BI, Tableau, DAX, dbt, Google BigQuery, MySQL, SQL Server, data pipelines, dashboard design')
skill('Libraries & Methods ','Pandas, NumPy, Matplotlib, Seaborn, scikit-learn, hypothesis testing, regression and classification')

heading('EXPERIENCE')
entry('Test Administrator / Data Support | Pearson Professional Center','Mar 2026 - Present')
bullet('Automated daily candidate-roster generation with XLOOKUP, INDEX/MATCH, nested logic, and dynamic arrays, reducing manual entry and preventing roster errors.')
bullet('Extract and prepare daily check-in data while validating schedules and candidate records to maintain accurate operations.')
bullet('Standardized reporting templates and documented repeatable data workflows for the administration team.')
entry('Marketing Analyst | Heyo Smart Technology','Sep 2023 - May 2024')
bullet('Analyzed market, competitor, and customer behavior using Google Analytics; translated real-time metrics into KPI reports and campaign decisions.')
bullet('Ran A/B tests and targeted campaigns that increased website visitors by 50% and generated 150+ clicks in seven days.')
entry('SEO Analyst | EveHR LLC','Jun 2023 - Nov 2023')
bullet('Analyzed keyword and market trends with SEMrush and Ahrefs, informing SEO, social, and email investment decisions.')
bullet('Increased website traffic 70% in six months and achieved 20+ top-five keyword rankings through data-driven optimization.')

heading('SELECTED ANALYTICS PROJECTS')
entry('Exoplanet Habitability Analysis | SQL, Python, Power BI, dbt, BigQuery','May 2025')
bullet('Built a cloud analytics pipeline from API extraction and Google Cloud Storage through tested dbt transformations in BigQuery.')
bullet('Designed an interactive Power BI dashboard to explore planetary characteristics and habitability trends.')
entry('Gym Check-ins & User Analysis | MySQL, Tableau','Dec 2024')
bullet('Queried 300K+ records across five related tables using joins, CTEs, subqueries, and stored functions; visualized usage and customer patterns in Tableau.')
bullet('Identified Gen Y as the most active segment, with particularly high female engagement, supporting targeted marketing opportunities.')
entry('Real Estate Sales Analysis | Python, Pandas, Visualization','Jan 2025')
bullet('Cleaned and analyzed 1M+ property-sale records, resolving missing values, inconsistencies, and outliers before exploratory analysis.')
bullet('Found that homes sold for roughly 1.6x assessed value and that 75% took more than one year to sell after listing.')
entry('Crash Reporting Dashboard | Tableau, Power BI, DAX','Nov 2024')
bullet('Collaborated on dashboards for 18K+ collision records using time-series and geographic analysis, calculated fields, table calculations, and DAX.')

heading('EDUCATION')
entry('Diploma, Computer Studies & Information Systems (Data Analytics) | Douglas College','Expected Sep 2026')
entry('Bachelor of Business Management with Marketing | University of the West of England','Sep 2023')
textline('Leadership: Vice President, Flagship Startup Club',italic=True,left=-.5)

doc.core_properties.title='Demi Le - Data Analyst Resume'
doc.core_properties.subject='Resume formatted using supplied data analyst template'
doc.core_properties.author='Demi Le'
doc.save(OUTPUT)
assert hashlib.sha256(open(REFERENCE,'rb').read()).hexdigest()==EXPECTED
print(OUTPUT)
